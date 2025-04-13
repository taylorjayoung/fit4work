"""
Resume Generator for the Job Scraper Application.

This module provides functionality for generating tailored resumes
based on job listings and user resumes.
"""

import os
import re
import logging
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import spacy
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords

logger = logging.getLogger(__name__)

# Ensure NLTK resources are downloaded
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_md")
except OSError:
    logger.warning("Downloading spaCy model. This may take a while...")
    spacy.cli.download("en_core_web_md")
    nlp = spacy.load("en_core_web_md")

class ResumeGenerator:
    """
    Generator for creating tailored resumes.
    
    This class provides methods for analyzing job listings and user resumes,
    and generating tailored resumes based on the analysis.
    """
    
    def __init__(self, output_dir):
        """
        Initialize the resume generator.
        
        Args:
            output_dir: Directory to save generated resumes
        """
        self.output_dir = Path(output_dir)
        self.stop_words = set(stopwords.words('english'))
        
        # Create output directory if it doesn't exist
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate_tailored_resume(self, resume_data, job_listing, output_filename=None):
        """
        Generate a tailored resume based on a job listing.
        
        Args:
            resume_data: Dictionary containing resume information
            job_listing: Dictionary containing job listing information
            output_filename: Name of the output file (if None, generate a name)
            
        Returns:
            Path to the generated resume file
        """
        try:
            # Extract keywords from the job listing
            job_keywords = self._extract_keywords_from_job(job_listing)
            
            # Match resume skills with job keywords
            skill_matches = self._match_skills_with_keywords(resume_data['skills'], job_keywords)
            
            # Generate a tailored resume
            if not output_filename:
                company_name = re.sub(r'[^\w\s-]', '', job_listing['company_name']).strip()
                job_title = re.sub(r'[^\w\s-]', '', job_listing['title']).strip()
                output_filename = f"Tailored_Resume_{company_name}_{job_title}.docx"
            
            output_path = self.output_dir / output_filename
            
            # Create the tailored resume document
            doc = self._create_resume_document(resume_data, job_listing, skill_matches, job_keywords)
            
            # Save the document
            doc.save(output_path)
            
            logger.info(f"Generated tailored resume: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error generating tailored resume: {e}", exc_info=True)
            return None
    
    def _extract_keywords_from_job(self, job_listing):
        """
        Extract keywords from a job listing.
        
        Args:
            job_listing: Dictionary containing job listing information
            
        Returns:
            List of extracted keywords
        """
        keywords = []
        
        # Extract text from the job description
        description = job_listing.get('description', '')
        if not description:
            return keywords
        
        # Process the text with spaCy
        doc = nlp(description)
        
        # Extract nouns and noun phrases
        for chunk in doc.noun_chunks:
            if chunk.root.pos_ in ['NOUN', 'PROPN'] and chunk.root.is_alpha:
                keyword = chunk.text.lower()
                if keyword not in self.stop_words and len(keyword) > 1:
                    keywords.append(keyword)
        
        # Extract named entities
        for ent in doc.ents:
            if ent.label_ in ['ORG', 'PRODUCT', 'WORK_OF_ART', 'LAW', 'LANGUAGE']:
                keyword = ent.text.lower()
                if keyword not in self.stop_words and len(keyword) > 1:
                    keywords.append(keyword)
        
        # Extract technical skills and buzzwords
        skill_patterns = [
            r'\b(?:proficient|experienced|skilled|expertise|knowledge)\s+(?:in|with)?\s+([A-Za-z0-9\+\#\s]+)',
            r'\b(?:familiarity|familiar)\s+(?:with)?\s+([A-Za-z0-9\+\#\s]+)',
            r'\b(?:experience|background)\s+(?:in|with)?\s+([A-Za-z0-9\+\#\s]+)',
            r'\b(?:understanding|proficiency)\s+(?:of|in|with)?\s+([A-Za-z0-9\+\#\s]+)'
        ]
        
        for pattern in skill_patterns:
            matches = re.findall(pattern, description, re.IGNORECASE)
            for match in matches:
                keyword = match.lower().strip()
                if keyword not in self.stop_words and len(keyword) > 1:
                    keywords.append(keyword)
        
        # Look for specific technical skills
        tech_skills = [
            'python', 'java', 'javascript', 'html', 'css', 'sql', 'nosql', 'react', 'angular', 'vue',
            'node', 'express', 'django', 'flask', 'spring', 'aws', 'azure', 'gcp', 'docker', 'kubernetes',
            'git', 'agile', 'scrum', 'jira', 'confluence', 'excel', 'word', 'powerpoint', 'photoshop',
            'illustrator', 'indesign', 'figma', 'sketch', 'leadership', 'management', 'communication',
            'teamwork', 'problem-solving', 'analytical', 'critical thinking', 'time management',
            'project management', 'research', 'writing', 'editing', 'public speaking', 'customer service',
            'sales', 'marketing', 'finance', 'accounting', 'hr', 'recruiting', 'training', 'coaching',
            'mentoring', 'data analysis', 'data visualization', 'machine learning', 'ai', 'nlp',
            'computer vision', 'statistics', 'calculus', 'linear algebra', 'probability'
        ]
        
        for skill in tech_skills:
            if re.search(r'\b' + re.escape(skill) + r'\b', description, re.IGNORECASE):
                keywords.append(skill)
        
        # Remove duplicates and sort by length (longer phrases first)
        unique_keywords = list(set(keywords))
        unique_keywords.sort(key=len, reverse=True)
        
        return unique_keywords
    
    def _match_skills_with_keywords(self, skills, keywords):
        """
        Match resume skills with job keywords.
        
        Args:
            skills: List of skills from the resume
            keywords: List of keywords from the job listing
            
        Returns:
            Dictionary mapping skills to matching keywords and scores
        """
        matches = {}
        
        if not skills or not keywords:
            return matches
        
        # Convert skills and keywords to spaCy docs
        skill_docs = [nlp(skill.lower()) for skill in skills]
        keyword_docs = [nlp(keyword.lower()) for keyword in keywords]
        
        # Match skills with keywords
        for i, skill_doc in enumerate(skill_docs):
            skill = skills[i]
            best_match = None
            best_score = 0
            
            for keyword_doc in keyword_docs:
                # Calculate similarity score
                score = skill_doc.similarity(keyword_doc)
                
                if score > best_score and score > 0.6:  # Threshold for considering a match
                    best_score = score
                    best_match = keyword_doc.text
            
            if best_match:
                matches[skill] = {
                    'keyword': best_match,
                    'score': best_score
                }
        
        return matches
    
    def _create_resume_document(self, resume_data, job_listing, skill_matches, job_keywords):
        """
        Create a tailored resume document.
        
        Args:
            resume_data: Dictionary containing resume information
            job_listing: Dictionary containing job listing information
            skill_matches: Dictionary mapping skills to matching keywords and scores
            job_keywords: List of keywords from the job listing
            
        Returns:
            docx.Document object containing the tailored resume
        """
        doc = docx.Document()
        
        # Set up document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Add name
        name = resume_data.get('name', 'Your Name')
        name_paragraph = doc.add_paragraph()
        name_run = name_paragraph.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(16)
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contact information
        contact_info = []
        
        # Add email
        emails = resume_data.get('email', [])
        if emails:
            contact_info.append(emails[0])
        
        # Add phone
        phones = resume_data.get('phone', [])
        if phones:
            contact_info.append(phones[0])
        
        # Add links
        links = resume_data.get('links', {})
        for link_type, url in links.items():
            if link_type in ['linkedin', 'github', 'portfolio']:
                contact_info.append(url)
        
        # Add contact information to document
        contact_paragraph = doc.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_paragraph.add_run(' | '.join(contact_info))
        
        # Add a horizontal line
        doc.add_paragraph('_' * 80)
        
        # Add professional summary
        doc.add_heading('Professional Summary', level=1)
        summary = self._generate_professional_summary(resume_data, job_listing, job_keywords)
        doc.add_paragraph(summary)
        
        # Add skills section (prioritizing matching skills)
        doc.add_heading('Skills', level=1)
        skills_paragraph = doc.add_paragraph()
        
        # First add skills that match job keywords
        matching_skills = [skill for skill in resume_data.get('skills', []) if skill in skill_matches]
        other_skills = [skill for skill in resume_data.get('skills', []) if skill not in skill_matches]
        
        all_skills = matching_skills + other_skills
        if all_skills:
            skills_paragraph.add_run(', '.join(all_skills))
        else:
            skills_paragraph.add_run('No skills listed')
        
        # Add experience section
        doc.add_heading('Experience', level=1)
        experiences = resume_data.get('experience', [])
        if experiences:
            for experience in experiences:
                # Highlight experience that contains job keywords
                highlighted_experience = self._highlight_keywords_in_text(experience, job_keywords)
                doc.add_paragraph(highlighted_experience)
        else:
            doc.add_paragraph('No experience listed')
        
        # Add education section
        doc.add_heading('Education', level=1)
        education = resume_data.get('education', [])
        if education:
            for edu in education:
                doc.add_paragraph(edu)
        else:
            doc.add_paragraph('No education listed')
        
        return doc
    
    def _generate_professional_summary(self, resume_data, job_listing, job_keywords):
        """
        Generate a professional summary tailored to the job listing.
        
        Args:
            resume_data: Dictionary containing resume information
            job_listing: Dictionary containing job listing information
            job_keywords: List of keywords from the job listing
            
        Returns:
            Generated professional summary
        """
        # Extract job title and company name
        job_title = job_listing.get('title', 'the position')
        company_name = job_listing.get('company_name', 'your company')
        
        # Extract years of experience from resume
        years_of_experience = self._extract_years_of_experience(resume_data)
        
        # Extract top skills that match job keywords
        skills = resume_data.get('skills', [])
        matching_skills = [skill for skill in skills if any(keyword in skill.lower() for keyword in job_keywords)]
        top_skills = matching_skills[:3] if matching_skills else skills[:3]
        
        # Generate the summary
        summary = f"Experienced professional with {years_of_experience}+ years of experience, seeking the {job_title} role at {company_name}. "
        
        if top_skills:
            summary += f"Skilled in {', '.join(top_skills[:-1])}"
            if len(top_skills) > 1:
                summary += f", and {top_skills[-1]}"
            elif len(top_skills) == 1:
                summary += f"{top_skills[0]}"
            summary += ". "
        
        summary += "Dedicated to delivering high-quality results, with a proven track record of success in fast-paced environments. "
        summary += "Eager to contribute expertise and skills to help achieve company goals and drive success."
        
        return summary
    
    def _extract_years_of_experience(self, resume_data):
        """
        Extract years of experience from resume data.
        
        Args:
            resume_data: Dictionary containing resume information
            
        Returns:
            Estimated years of experience
        """
        experiences = resume_data.get('experience', [])
        
        if not experiences:
            return 2  # Default value if no experience is listed
        
        # Look for year ranges in experience entries
        years = []
        year_pattern = r'\b(19|20)\d{2}\b'
        
        for experience in experiences:
            matches = re.findall(year_pattern, experience)
            if matches:
                years.extend([int(year) for year in matches])
        
        if len(years) >= 2:
            # Calculate the difference between the earliest and latest years
            years.sort()
            return max(1, years[-1] - years[0])
        
        # If we can't extract years, estimate based on the number of experience entries
        return max(1, len(experiences))
    
    def _highlight_keywords_in_text(self, text, keywords):
        """
        Highlight job keywords in text by adding asterisks around them.
        
        Args:
            text: Text to highlight keywords in
            keywords: List of keywords to highlight
            
        Returns:
            Text with highlighted keywords
        """
        highlighted_text = text
        
        for keyword in keywords:
            # Escape special regex characters in the keyword
            escaped_keyword = re.escape(keyword)
            
            # Replace the keyword with a highlighted version (surrounded by asterisks)
            pattern = r'\b' + escaped_keyword + r'\b'
            replacement = r'*\g<0>*'
            highlighted_text = re.sub(pattern, replacement, highlighted_text, flags=re.IGNORECASE)
        
        return highlighted_text
