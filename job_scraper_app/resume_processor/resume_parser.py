"""
Resume Parser for the Job Scraper Application.

This module provides functionality for parsing and extracting information
from user resumes in various formats (DOCX, PDF) using both traditional methods
and AI-powered parsing via Anthropic's Claude API.
"""

import os
import re
import json
import logging
import sys
from pathlib import Path
import docx
from pdfminer.high_level import extract_text
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from anthropic import Anthropic

# Add the parent directory to the Python path to import config_loader
parent_dir = str(Path(__file__).resolve().parent.parent)
if parent_dir not in sys.path:
    sys.path.insert(0, parent_dir)
from config_loader import get_anthropic_api_key

logger = logging.getLogger(__name__)

# Ensure NLTK resources are downloaded
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

class ResumeParser:
    """
    Parser for extracting information from resumes.
    
    This class provides methods for parsing resumes in different formats
    and extracting structured information from them using both traditional
    pattern matching and AI-powered parsing.
    """
    
    def __init__(self, config=None):
        """
        Initialize the resume parser.
        
        Args:
            config: Configuration dictionary containing AI service settings
        """
        self.stop_words = set(stopwords.words('english'))
        self.config = config
        self.anthropic_client = None
        
        # Try to get API key from config_loader if not provided in config
        api_key = None
        
        # Initialize Anthropic client if configured
        if config and 'ai_services' in config and 'anthropic' in config['ai_services']:
            anthropic_config = config['ai_services']['anthropic']
            if anthropic_config.get('enabled', False):
                api_key = anthropic_config.get('api_key')
                
        # If API key not found in config, try to get it from config_loader
        if not api_key or api_key == 'YOUR_ANTHROPIC_API_KEY':
            try:
                api_key = get_anthropic_api_key()
                logger.info("Using API key from environment or config_loader")
                
                # Update config with the new API key
                if self.config is None:
                    self.config = {}
                if 'ai_services' not in self.config:
                    self.config['ai_services'] = {}
                if 'anthropic' not in self.config['ai_services']:
                    self.config['ai_services']['anthropic'] = {'enabled': True}
                self.config['ai_services']['anthropic']['api_key'] = api_key
            except Exception as e:
                logger.error(f"Failed to get API key from config_loader: {e}", exc_info=True)
        
        # Initialize Anthropic client with the API key
        if api_key and api_key != 'YOUR_ANTHROPIC_API_KEY':
            try:
                self.anthropic_client = Anthropic(api_key=api_key)
                logger.info("Anthropic client initialized successfully")
            except Exception as e:
                logger.error(f"Failed to initialize Anthropic client: {e}", exc_info=True)
    
    def parse(self, file_path):
        """
        Parse a resume file and extract structured information.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dictionary containing extracted resume information
        """
        try:
            logger.info(f"Starting to parse resume: {file_path}")
            logger.info(f"Anthropic client initialized: {self.anthropic_client is not None}")
            
            if self.config and 'ai_services' in self.config and 'anthropic' in self.config['ai_services']:
                api_key = self.config['ai_services']['anthropic'].get('api_key', '')
                logger.info(f"API key configured: {api_key[:5]}...{api_key[-4:] if len(api_key) > 10 else ''}")
            else:
                logger.warning("No Anthropic API configuration found in config")
            
            file_path = Path(file_path)
            
            # Check if the file exists
            if not file_path.exists():
                logger.error(f"Resume file not found: {file_path}")
                return None
            
            # Extract text based on file format
            if file_path.suffix.lower() == '.docx':
                text = self._extract_text_from_docx(file_path)
            elif file_path.suffix.lower() == '.pdf':
                text = self._extract_text_from_pdf(file_path)
            elif file_path.suffix.lower() == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                logger.error(f"Unsupported file format: {file_path.suffix}")
                return None
            
            if not text:
                logger.error(f"Failed to extract text from resume: {file_path}")
                return None
            
            # Try AI-powered parsing first if available
            if self.anthropic_client:
                try:
                    logger.info(f"Attempting AI-powered parsing for resume: {file_path}")
                    logger.info(f"Anthropic client initialized: {self.anthropic_client is not None}")
                    
                    # Check if API key is set
                    if self.config and 'ai_services' in self.config and 'anthropic' in self.config['ai_services']:
                        api_key = self.config['ai_services']['anthropic'].get('api_key', '')
                        logger.info(f"API key configured: {api_key[:5]}...{api_key[-5:] if len(api_key) > 10 else ''}")
                    
                    # Try to parse with Anthropic
                    resume_data = self._parse_with_anthropic(text)
                    
                    if resume_data:
                        logger.info(f"AI parsing successful, got data: {list(resume_data.keys())}")
                        resume_data['file_path'] = str(file_path)
                        resume_data['content_text'] = text
                        logger.info(f"Successfully parsed resume with AI: {file_path}")
                        return resume_data
                    else:
                        logger.warning("AI parsing returned None, falling back to traditional parsing")
                except Exception as e:
                    logger.warning(f"AI-powered parsing failed, falling back to traditional parsing: {e}", exc_info=True)
            else:
                logger.info("Anthropic client not initialized, using traditional parsing")
            
            # Fall back to traditional parsing methods
            logger.info("Using traditional parsing methods")
            resume_data = self._extract_information(text)
            resume_data['file_path'] = str(file_path)
            resume_data['content_text'] = text
            
            logger.info(f"Successfully parsed resume: {file_path}")
            return resume_data
            
        except Exception as e:
            logger.error(f"Error parsing resume {file_path}: {e}", exc_info=True)
            return None
            
    def _parse_with_anthropic(self, text):
        """
        Parse resume text using Anthropic's Claude API.
        
        Args:
            text: Resume text to parse
            
        Returns:
            Dictionary containing structured resume information
        """
        if not self.anthropic_client:
            logger.error("Anthropic client not initialized")
            return None
            
        # Get Anthropic configuration
        anthropic_config = self.config['ai_services']['anthropic']
        model = anthropic_config.get('model', 'claude-3-opus-20240229')
        max_tokens = anthropic_config.get('max_tokens', 4096)
        temperature = anthropic_config.get('temperature', 0.2)
        
        # Create the prompt
        prompt = f"""
        You are an expert resume parser. I will provide you with the text of a resume, and I need you to extract structured information from it.
        
        Please analyze the resume and extract the following sections:
        
        1. Personal Information:
           - Name of the candidate
           - Email addresses
           - Phone numbers
           - Location/address
           - Links (LinkedIn, GitHub, portfolio, etc.)
        
        2. Education:
           - Each entry should include institution, degree, field of study, and dates if available
           - Include relevant coursework if mentioned
        
        3. Professional Experience:
           - Each entry should include company name, job title, dates, and description
           - Include responsibilities, achievements, and metrics where available
        
        4. Projects:
           - Each project should include name, description, technologies used, and dates if available
           - Include links to projects if available
        
        5. Skills:
           - Technical skills (programming languages, tools, platforms, etc.)
           - Soft skills (communication, leadership, etc.)
           - Group skills by category if possible
        
        6. Volunteer Work:
           - Organization name, role, dates, and description
           - Include responsibilities and achievements
        
        7. Miscellaneous:
           - Any other relevant information such as certifications, awards, languages, interests, etc.
           - Include anything that doesn't fit in the above categories
        
        Format your response as a JSON object with the following structure:
        {{
            "personal_info": {{
                "name": "Candidate Name",
                "email": ["email1@example.com", "email2@example.com"],
                "phone": ["123-456-7890"],
                "location": "City, State, Country",
                "links": {{
                    "linkedin": "LinkedIn URL",
                    "github": "GitHub URL",
                    "portfolio": "Portfolio URL",
                    "other": ["Other URL 1", "Other URL 2"]
                }}
            }},
            "education": [
                {{
                    "institution": "University Name",
                    "degree": "Degree Name",
                    "field": "Field of Study",
                    "start_date": "Start Date",
                    "end_date": "End Date",
                    "coursework": ["Course 1", "Course 2"],
                    "gpa": "GPA if available"
                }}
            ],
            "professional_experience": [
                {{
                    "company": "Company Name",
                    "title": "Job Title",
                    "start_date": "Start Date",
                    "end_date": "End Date",
                    "description": "Job Description",
                    "responsibilities": ["Responsibility 1", "Responsibility 2"],
                    "achievements": ["Achievement 1", "Achievement 2"]
                }}
            ],
            "projects": [
                {{
                    "name": "Project Name",
                    "description": "Project Description",
                    "technologies": ["Tech 1", "Tech 2"],
                    "start_date": "Start Date",
                    "end_date": "End Date",
                    "link": "Project URL if available"
                }}
            ],
            "skills": {{
                "technical": ["Skill 1", "Skill 2", "Skill 3"],
                "soft": ["Skill 1", "Skill 2", "Skill 3"],
                "languages": ["Language 1", "Language 2"],
                "tools": ["Tool 1", "Tool 2"]
            }},
            "volunteer_work": [
                {{
                    "organization": "Organization Name",
                    "role": "Role Title",
                    "start_date": "Start Date",
                    "end_date": "End Date",
                    "description": "Description of volunteer work",
                    "achievements": ["Achievement 1", "Achievement 2"]
                }}
            ],
            "miscellaneous": {{
                "certifications": ["Certification 1", "Certification 2"],
                "awards": ["Award 1", "Award 2"],
                "languages": ["Language 1 (Proficiency)", "Language 2 (Proficiency)"],
                "interests": ["Interest 1", "Interest 2"],
                "other": "Any other relevant information"
            }}
        }}
        
        If any information is not available in the resume, use null or an empty array/object as appropriate.
        
        Here is the resume text:
        
        {text}
        """
        
        try:
            logger.info(f"Calling Anthropic API with model: {model}, max_tokens: {max_tokens}, temperature: {temperature}")
            
            # Call the Anthropic API
            response = self.anthropic_client.messages.create(
                model=model,
                max_tokens=max_tokens,
                temperature=temperature,
                system="You are an expert resume parser that extracts structured information from resumes and returns it in JSON format.",
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            logger.info(f"Received response from Anthropic API: {response.id}")
            
            # Extract the JSON response
            response_text = response.content[0].text
            logger.info(f"Response text length: {len(response_text)} characters")
            
            # Find the JSON part of the response
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1
            logger.info(f"JSON part found at positions {json_start} to {json_end}")
            
            if json_start >= 0 and json_end > json_start:
                json_str = response_text[json_start:json_end]
                resume_data = json.loads(json_str)
                
                # Convert to the format expected by the application
                formatted_data = {
                    'name': resume_data.get('personal_info', {}).get('name'),
                    'email': resume_data.get('personal_info', {}).get('email', []),
                    'phone': resume_data.get('personal_info', {}).get('phone', []),
                    'location': resume_data.get('personal_info', {}).get('location'),
                    'education': [],
                    'experience': [],
                    'skills': [],
                    'links': resume_data.get('personal_info', {}).get('links', {})
                }
                
                # Format education entries
                for edu in resume_data.get('education', []):
                    entry = f"{edu.get('institution', '')} - {edu.get('degree', '')} {edu.get('field', '')}"
                    if edu.get('start_date') and edu.get('end_date'):
                        entry += f" ({edu.get('start_date')} - {edu.get('end_date')})"
                    if edu.get('gpa'):
                        entry += f", GPA: {edu.get('gpa')}"
                    if edu.get('coursework') and len(edu.get('coursework')) > 0:
                        entry += f"\nCoursework: {', '.join(edu.get('coursework'))}"
                    formatted_data['education'].append(entry)
                
                # Format professional experience entries
                for exp in resume_data.get('professional_experience', []):
                    entry = f"{exp.get('title', '')} at {exp.get('company', '')}"
                    if exp.get('start_date') and exp.get('end_date'):
                        entry += f" ({exp.get('start_date')} - {exp.get('end_date')})"
                    if exp.get('description'):
                        entry += f"\n{exp.get('description')}"
                    
                    # Add responsibilities
                    if exp.get('responsibilities') and len(exp.get('responsibilities')) > 0:
                        entry += "\nResponsibilities:"
                        for resp in exp.get('responsibilities'):
                            entry += f"\n- {resp}"
                    
                    # Add achievements
                    if exp.get('achievements') and len(exp.get('achievements')) > 0:
                        entry += "\nAchievements:"
                        for ach in exp.get('achievements'):
                            entry += f"\n- {ach}"
                            
                    formatted_data['experience'].append(entry)
                
                # Add projects as additional experience entries
                for proj in resume_data.get('projects', []):
                    entry = f"Project: {proj.get('name', '')}"
                    if proj.get('start_date') and proj.get('end_date'):
                        entry += f" ({proj.get('start_date')} - {proj.get('end_date')})"
                    if proj.get('description'):
                        entry += f"\n{proj.get('description')}"
                    if proj.get('technologies') and len(proj.get('technologies')) > 0:
                        entry += f"\nTechnologies: {', '.join(proj.get('technologies'))}"
                    if proj.get('link'):
                        entry += f"\nLink: {proj.get('link')}"
                    formatted_data['experience'].append(entry)
                
                # Add volunteer work as additional experience entries
                for vol in resume_data.get('volunteer_work', []):
                    entry = f"Volunteer: {vol.get('role', '')} at {vol.get('organization', '')}"
                    if vol.get('start_date') and vol.get('end_date'):
                        entry += f" ({vol.get('start_date')} - {vol.get('end_date')})"
                    if vol.get('description'):
                        entry += f"\n{vol.get('description')}"
                    if vol.get('achievements') and len(vol.get('achievements')) > 0:
                        entry += "\nAchievements:"
                        for ach in vol.get('achievements'):
                            entry += f"\n- {ach}"
                    formatted_data['experience'].append(entry)
                
                # Process skills
                skills = resume_data.get('skills', {})
                if isinstance(skills, dict):
                    # Technical skills
                    if skills.get('technical'):
                        formatted_data['skills'].extend(skills.get('technical', []))
                    
                    # Soft skills
                    if skills.get('soft'):
                        formatted_data['skills'].extend(skills.get('soft', []))
                    
                    # Languages
                    if skills.get('languages'):
                        formatted_data['skills'].extend(skills.get('languages', []))
                    
                    # Tools
                    if skills.get('tools'):
                        formatted_data['skills'].extend(skills.get('tools', []))
                elif isinstance(skills, list):
                    # If skills is a list, just add them directly
                    formatted_data['skills'].extend(skills)
                
                # Add miscellaneous information
                misc = resume_data.get('miscellaneous', {})
                if misc:
                    misc_entry = "Additional Information:"
                    
                    # Certifications
                    if misc.get('certifications') and len(misc.get('certifications')) > 0:
                        misc_entry += f"\nCertifications: {', '.join(misc.get('certifications'))}"
                    
                    # Awards
                    if misc.get('awards') and len(misc.get('awards')) > 0:
                        misc_entry += f"\nAwards: {', '.join(misc.get('awards'))}"
                    
                    # Languages
                    if misc.get('languages') and len(misc.get('languages')) > 0:
                        misc_entry += f"\nLanguages: {', '.join(misc.get('languages'))}"
                    
                    # Interests
                    if misc.get('interests') and len(misc.get('interests')) > 0:
                        misc_entry += f"\nInterests: {', '.join(misc.get('interests'))}"
                    
                    # Other
                    if misc.get('other'):
                        misc_entry += f"\nOther: {misc.get('other')}"
                    
                    if misc_entry != "Additional Information:":
                        formatted_data['experience'].append(misc_entry)
                
                return formatted_data
            else:
                logger.error("Failed to extract JSON from Anthropic response")
                return None
                
        except Exception as e:
            logger.error(f"Error calling Anthropic API: {e}", exc_info=True)
            return None
    
    def _extract_text_from_docx(self, file_path):
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text as a string
        """
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX file {file_path}: {e}", exc_info=True)
            return None
    
    def _extract_text_from_pdf(self, file_path):
        """
        Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text as a string
        """
        try:
            return extract_text(file_path)
        except Exception as e:
            logger.error(f"Error extracting text from PDF file {file_path}: {e}", exc_info=True)
            return None
    
    def _extract_information(self, text):
        """
        Extract structured information from resume text.
        
        Args:
            text: Resume text
            
        Returns:
            Dictionary containing extracted information
        """
        resume_data = {
            'name': self._extract_name(text),
            'email': self._extract_email(text),
            'phone': self._extract_phone(text),
            'education': self._extract_education(text),
            'experience': self._extract_experience(text),
            'skills': self._extract_skills(text),
            'links': self._extract_links(text)
        }
        
        return resume_data
    
    def _extract_name(self, text):
        """
        Extract the candidate's name from the resume text.
        
        Args:
            text: Resume text
            
        Returns:
            Extracted name or None if not found
        """
        # Typically, the name is at the beginning of the resume
        lines = text.split('\n')
        non_empty_lines = [line.strip() for line in lines if line.strip()]
        
        if non_empty_lines:
            # Assume the first non-empty line is the name
            # This is a simple heuristic and may need to be improved
            return non_empty_lines[0]
        
        return None
    
    def _extract_email(self, text):
        """
        Extract email addresses from the resume text.
        
        Args:
            text: Resume text
            
        Returns:
            List of extracted email addresses
        """
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        return re.findall(email_pattern, text)
    
    def _extract_phone(self, text):
        """
        Extract phone numbers from the resume text.
        
        Args:
            text: Resume text
            
        Returns:
            List of extracted phone numbers
        """
        phone_pattern = r'\b(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b'
        return re.findall(phone_pattern, text)
    
    def _extract_education(self, text):
        """
        Extract education information from the resume text.
        
        Args:
            text: Resume text
            
        Returns:
            List of extracted education entries
        """
        education = []
        
        # Look for education section
        education_section = self._extract_section(text, ['education', 'academic background', 'academic history'])
        
        if education_section:
            # Look for degree keywords and institution names
            degree_keywords = [
                'bachelor', 'master', 'phd', 'doctorate', 'bs', 'ba', 'ms', 'ma', 'mba',
                'associate', 'certificate', 'certification', 'diploma', 'degree', 'b.s.', 'b.a.',
                'm.s.', 'm.a.', 'ph.d.', 'b.sc', 'm.sc', 'bsc', 'msc', 'btech', 'mtech'
            ]
            
            institution_keywords = [
                'university', 'college', 'institute', 'school', 'academy', 'polytechnic'
            ]
            
            # Look for year patterns (e.g., 2015-2019, 2020 - Present)
            year_pattern = r'(?:19|20)\d{2}(?:\s*[-–—]\s*(?:(?:19|20)\d{2}|present|current|now))?'
            
            # Split into paragraphs (each education entry might be a paragraph)
            paragraphs = re.split(r'\n\s*\n', education_section)
            
            for paragraph in paragraphs:
                if paragraph.strip():
                    # Check if this paragraph contains education information
                    has_degree = any(keyword in paragraph.lower() for keyword in degree_keywords)
                    has_institution = any(keyword in paragraph.lower() for keyword in institution_keywords)
                    has_year = re.search(year_pattern, paragraph, re.IGNORECASE) is not None
                    
                    if has_degree or has_institution or has_year:
                        education.append(paragraph.strip())
            
            # If no paragraphs were found, try line by line
            if not education:
                lines = education_section.split('\n')
                i = 0
                while i < len(lines):
                    line = lines[i].strip()
                    if not line:
                        i += 1
                        continue
                    
                    # Check if this line contains education information
                    has_degree = any(keyword in line.lower() for keyword in degree_keywords)
                    has_institution = any(keyword in line.lower() for keyword in institution_keywords)
                    has_year = re.search(year_pattern, line, re.IGNORECASE) is not None
                    
                    if has_degree or has_institution or has_year:
                        # Include this line and possibly the next few lines as one education entry
                        entry = line
                        for j in range(1, 3):  # Look at the next 2 lines
                            if i + j < len(lines) and lines[i + j].strip():
                                entry += ' ' + lines[i + j].strip()
                        
                        education.append(entry.strip())
                        i += 3  # Skip the lines we've included
                    else:
                        i += 1
        
        # If still no education entries found, try to find education information in the entire text
        if not education:
            # Look for patterns like "Bachelor of Science in Computer Science"
            degree_patterns = [
                r'(?:bachelor|master|doctor|associate)(?:\s+of\s+)(?:science|arts|business|engineering|fine arts|education|law|medicine|nursing|philosophy|psychology|social work|technology)',
                r'(?:bs|ba|ms|ma|mba|phd|md|jd)(?:\s+in\s+)(?:\w+(?:\s+\w+)*)',
                r'(?:bachelor|master|doctor|associate)(?:\'s)?(?:\s+degree)?(?:\s+in\s+)(?:\w+(?:\s+\w+)*)'
            ]
            
            for pattern in degree_patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    # Get the sentence containing the match
                    start = text.rfind('.', 0, match.start()) + 1
                    end = text.find('.', match.end())
                    if end == -1:
                        end = len(text)
                    
                    sentence = text[start:end].strip()
                    if sentence and sentence not in education:
                        education.append(sentence)
        
        return education
    
    def _extract_experience(self, text):
        """
        Extract work experience information from the resume text.
        
        Args:
            text: Resume text
            
        Returns:
            List of extracted experience entries
        """
        experience = []
        
        # Look for experience section
        experience_section = self._extract_section(text, ['experience', 'work experience', 'employment', 'work history'])
        
        if experience_section:
            # Look for company names, job titles, and dates
            company_keywords = [
                'inc', 'llc', 'corporation', 'corp', 'company', 'co', 'ltd', 'limited',
                'group', 'technologies', 'solutions', 'systems', 'services', 'associates'
            ]
            
            job_title_keywords = [
                'engineer', 'developer', 'manager', 'director', 'analyst', 'specialist',
                'consultant', 'coordinator', 'assistant', 'associate', 'lead', 'senior',
                'junior', 'intern', 'administrator', 'supervisor', 'head', 'chief',
                'officer', 'president', 'vice president', 'vp', 'ceo', 'cto', 'cfo', 'coo'
            ]
            
            # Look for date patterns (e.g., Jan 2020 - Present, 2018-2020)
            date_patterns = [
                r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}\s*[-–—]\s*(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}',
                r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}\s*[-–—]\s*(?:present|current|now)',
                r'(?:19|20)\d{2}\s*[-–—]\s*(?:19|20)\d{2}',
                r'(?:19|20)\d{2}\s*[-–—]\s*(?:present|current|now)'
            ]
            
            # Try to identify job entries by looking for patterns
            # First, split by potential job entry markers (dates, company names, etc.)
            potential_markers = []
            
            # Add date markers
            for pattern in date_patterns:
                for match in re.finditer(pattern, experience_section, re.IGNORECASE):
                    potential_markers.append((match.start(), match.group()))
            
            # Add company name markers (lines containing company keywords that look like headers)
            lines = experience_section.split('\n')
            current_pos = 0
            for line in lines:
                line_lower = line.lower().strip()
                if (len(line_lower) < 50 and 
                    any(keyword in line_lower for keyword in company_keywords) and
                    (line.isupper() or line.istitle() or line.endswith(':'))):
                    potential_markers.append((current_pos, line.strip()))
                current_pos += len(line) + 1  # +1 for the newline
            
            # Add job title markers
            for keyword in job_title_keywords:
                for match in re.finditer(r'\b{}\b'.format(re.escape(keyword)), experience_section, re.IGNORECASE):
                    # Get the line containing this keyword
                    line_start = experience_section.rfind('\n', 0, match.start()) + 1
                    line_end = experience_section.find('\n', match.end())
                    if line_end == -1:
                        line_end = len(experience_section)
                    
                    line = experience_section[line_start:line_end].strip()
                    # Check if this line looks like a job title
                    if (len(line) < 50 and 
                        (line.istitle() or line.isupper() or line.endswith(':')) and
                        not any(marker[1] == line for marker in potential_markers)):
                        potential_markers.append((line_start, line))
            
            # Sort markers by position
            potential_markers.sort(key=lambda x: x[0])
            
            # Extract job entries based on markers
            if potential_markers:
                for i in range(len(potential_markers)):
                    start_pos = potential_markers[i][0]
                    end_pos = potential_markers[i+1][0] if i < len(potential_markers) - 1 else len(experience_section)
                    
                    # Extract the job entry
                    job_entry = experience_section[start_pos:end_pos].strip()
                    if job_entry and len(job_entry) > 20:  # Minimum length to be considered a valid entry
                        experience.append(job_entry)
            else:
                # If no markers found, fall back to paragraph splitting
                paragraphs = re.split(r'\n\s*\n', experience_section)
                for paragraph in paragraphs:
                    if paragraph.strip():
                        experience.append(paragraph.strip())
        
        # If no experience entries found, try to find experience information in the entire text
        if not experience:
            # Look for job title patterns
            job_title_patterns = [
                r'(?:senior|junior|lead|principal|staff|associate)?\s*(?:software|systems|data|web|frontend|backend|full-stack|fullstack|mobile|ios|android|devops|cloud|network|security|qa|test|database|machine learning|ai|ml|ui|ux)\s*(?:engineer|developer|architect|analyst|specialist|consultant)',
                r'(?:project|product|program|technical|engineering|development|team|group|department)\s*(?:manager|director|lead|head)',
                r'(?:chief|vice president|vp|director|head)\s*(?:technology|technical|information|product|executive|operating|financial|marketing|sales)'
            ]
            
            for pattern in job_title_patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    # Get the paragraph containing the match
                    para_start = text.rfind('\n\n', 0, match.start())
                    if para_start == -1:
                        para_start = 0
                    else:
                        para_start += 2  # Skip the newlines
                    
                    para_end = text.find('\n\n', match.end())
                    if para_end == -1:
                        para_end = len(text)
                    
                    paragraph = text[para_start:para_end].strip()
                    if paragraph and paragraph not in experience:
                        experience.append(paragraph)
        
        return experience
    
    def _extract_skills(self, text):
        """
        Extract skills from the resume text.
        
        Args:
            text: Resume text
            
        Returns:
            List of extracted skills
        """
        skills = []
        
        # Look for skills section
        skills_section = self._extract_section(text, ['skills', 'technical skills', 'core competencies', 'proficiencies'])
        
        if skills_section:
            # Tokenize and filter out stop words
            tokens = word_tokenize(skills_section.lower())
            filtered_tokens = [token for token in tokens if token.isalpha() and token not in self.stop_words]
            
            # Look for common skill keywords
            skill_keywords = [
                'python', 'java', 'javascript', 'html', 'css', 'sql', 'nosql', 'react', 'angular', 'vue',
                'node', 'express', 'django', 'flask', 'spring', 'aws', 'azure', 'gcp', 'docker', 'kubernetes',
                'git', 'agile', 'scrum', 'jira', 'confluence', 'excel', 'word', 'powerpoint', 'photoshop',
                'illustrator', 'indesign', 'figma', 'sketch', 'leadership', 'management', 'communication',
                'teamwork', 'problem-solving', 'analytical', 'critical thinking', 'time management',
                'project management', 'research', 'writing', 'editing', 'public speaking', 'customer service',
                'sales', 'marketing', 'finance', 'accounting', 'hr', 'recruiting', 'training', 'coaching',
                'mentoring', 'data analysis', 'data visualization', 'machine learning', 'ai', 'nlp',
                'computer vision', 'statistics', 'calculus', 'linear algebra', 'probability'
            ]
            
            # Extract skills from the skills section
            for keyword in skill_keywords:
                if keyword in ' '.join(filtered_tokens):
                    skills.append(keyword)
            
            # Also look for skills in bullet points
            bullet_points = re.findall(r'[•\-\*]\s*(.*?)(?:\n|$)', skills_section)
            for point in bullet_points:
                if point.strip():
                    skills.append(point.strip())
        
        return skills
    
    def _extract_links(self, text):
        """
        Extract links (e.g., LinkedIn, GitHub, portfolio) from the resume text.
        
        Args:
            text: Resume text
            
        Returns:
            Dictionary mapping link types to URLs
        """
        links = {}
        
        # Extract URLs
        url_pattern = r'https?://(?:www\.)?([A-Za-z0-9][-A-Za-z0-9]*\.)+[A-Za-z]{2,}(?:/[^\\s]*)?'
        urls = re.findall(url_pattern, text)
        
        # Categorize URLs
        for url in urls:
            if 'linkedin.com' in url.lower():
                links['linkedin'] = url
            elif 'github.com' in url.lower():
                links['github'] = url
            elif 'stackoverflow.com' in url.lower():
                links['stackoverflow'] = url
            elif any(domain in url.lower() for domain in ['portfolio', 'personal', 'website']):
                links['portfolio'] = url
            else:
                links.setdefault('other', []).append(url)
        
        return links
    
    def _extract_section(self, text, section_keywords):
        """
        Extract a section from the resume text based on section keywords.
        
        Args:
            text: Resume text
            section_keywords: List of keywords that might indicate the section
            
        Returns:
            Extracted section text or None if not found
        """
        # Convert text to lowercase for case-insensitive matching
        text_lower = text.lower()
        
        # Expanded section keywords with variations and synonyms
        expanded_keywords = []
        for keyword in section_keywords:
            expanded_keywords.append(keyword)
            # Add common variations
            if keyword == 'education':
                expanded_keywords.extend(['academic background', 'academic history', 'educational background', 
                                         'degrees', 'qualifications', 'academic qualifications', 'schooling',
                                         'academic experience', 'educational experience'])
            elif 'experience' in keyword:
                expanded_keywords.extend(['professional experience', 'work history', 'employment history', 
                                         'professional background', 'career history', 'job history',
                                         'professional summary', 'career summary', 'work summary'])
            elif keyword == 'skills':
                expanded_keywords.extend(['technical skills', 'core competencies', 'proficiencies', 
                                         'expertise', 'capabilities', 'qualifications', 'skill set',
                                         'technical proficiencies', 'areas of expertise'])
            elif keyword == 'projects':
                expanded_keywords.extend(['project experience', 'key projects', 'relevant projects',
                                         'personal projects', 'professional projects', 'project work'])
        
        # Find the start of the section using fuzzy matching
        section_start = -1
        best_match_score = 0
        
        # First try exact matches at the beginning of lines
        for keyword in expanded_keywords:
            # Look for the keyword at the beginning of a line
            pattern = r'(?:^|\n)(?:\s*)({})[:\s]'.format(re.escape(keyword))
            match = re.search(pattern, text_lower)
            if match:
                section_start = match.start()
                break
        
        # If no exact match, try fuzzy matching with section headers
        if section_start == -1:
            lines = text_lower.split('\n')
            for i, line in enumerate(lines):
                line = line.strip()
                # Check if this line looks like a section header (short, ends with colon, etc.)
                if (len(line) < 30 and (line.endswith(':') or line.isupper() or 
                                        any(char in line for char in ['━', '─', '=', '-', '_']))):
                    # Check similarity with our keywords
                    for keyword in expanded_keywords:
                        # Simple similarity: keyword is contained in the line
                        if keyword in line:
                            match_score = len(keyword) / len(line) if len(line) > 0 else 0
                            if match_score > best_match_score:
                                best_match_score = match_score
                                section_start = text_lower.find(line)
        
        if section_start == -1:
            return None
        
        # Find the end of the section (start of the next section)
        common_sections = [
            'education', 'experience', 'work experience', 'employment', 'skills',
            'technical skills', 'projects', 'publications', 'certifications',
            'awards', 'honors', 'languages', 'interests', 'references', 'summary',
            'objective', 'profile', 'contact', 'personal', 'professional',
            'qualifications', 'achievements', 'volunteer', 'activities'
        ]
        
        # Expanded common sections with variations
        expanded_common_sections = []
        for section in common_sections:
            expanded_common_sections.append(section)
            if section == 'education':
                expanded_common_sections.extend(['academic', 'degree', 'school', 'university', 'college'])
            elif section == 'experience' or section == 'work experience' or section == 'employment':
                expanded_common_sections.extend(['professional', 'career', 'job', 'work history'])
            elif section == 'skills' or section == 'technical skills':
                expanded_common_sections.extend(['competencies', 'proficiencies', 'expertise'])
        
        # Remove the current section from the list of common sections
        for keyword in expanded_keywords:
            if keyword in expanded_common_sections:
                expanded_common_sections.remove(keyword)
        
        section_end = len(text)
        
        # Look for the next section header
        lines = text_lower[section_start:].split('\n')
        current_pos = section_start
        for i, line in enumerate(lines):
            if i == 0:  # Skip the current section header
                current_pos += len(line) + 1  # +1 for the newline
                continue
                
            line = line.strip()
            # Check if this line looks like a section header
            if (len(line) < 30 and (line.endswith(':') or line.isupper() or 
                                    any(char in line for char in ['━', '─', '=', '-', '_']))):
                # Check if it matches any of our common section keywords
                for section in expanded_common_sections:
                    if section in line:
                        section_end = current_pos
                        break
                if section_end != len(text):
                    break
            
            current_pos += len(line) + 1  # +1 for the newline
        
        # Extract the section text
        section_text = text[section_start:section_end].strip()
        
        # Remove the section header
        lines = section_text.split('\n')
        if lines:
            section_text = '\n'.join(lines[1:])
        
        return section_text.strip()
