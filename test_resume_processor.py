#!/usr/bin/env python3
"""
Test script for the resume processor.

This script demonstrates how to test the resume processor by:
1. Creating a sample resume file
2. Uploading and parsing the resume
3. Generating a tailored resume for a job listing
4. Getting tailored resumes for a specific resume
"""

import os
import sys
import json
import logging
import tempfile
from pathlib import Path
from sqlalchemy import create_engine

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_sample_resume():
    """Create a sample resume file for testing."""
    try:
        # Import the required modules
        from docx import Document
        
        # Create a temporary directory for the sample resume
        temp_dir = tempfile.mkdtemp()
        resume_path = Path(temp_dir) / "sample_resume.docx"
        
        # Create a sample resume
        doc = Document()
        doc.add_heading('John Doe', 0)
        
        # Add contact information
        doc.add_paragraph('Email: john.doe@example.com')
        doc.add_paragraph('Phone: (123) 456-7890')
        doc.add_paragraph('LinkedIn: https://linkedin.com/in/johndoe')
        
        # Add education section
        doc.add_heading('Education', level=1)
        doc.add_paragraph('Bachelor of Science in Computer Science')
        doc.add_paragraph('University of Example, 2015-2019')
        
        # Add experience section
        doc.add_heading('Experience', level=1)
        doc.add_paragraph('Software Engineer at Example Corp, 2019-2021')
        doc.add_paragraph('Senior Software Engineer at Another Corp, 2021-Present')
        
        # Add skills section
        doc.add_heading('Skills', level=1)
        skills = doc.add_paragraph()
        skills.add_run('Python, Java, SQL, JavaScript, HTML, CSS, React, Node.js, Git, Docker')
        
        # Save the document
        doc.save(resume_path)
        
        logger.info(f"Created sample resume at {resume_path}")
        return resume_path
    
    except ImportError:
        logger.error("Failed to import docx module. Please install it with: pip install python-docx")
        return None
    
    except Exception as e:
        logger.error(f"Error creating sample resume: {e}")
        return None

def test_resume_processor():
    """Test the resume processor functionality."""
    try:
        # Create a sample resume
        resume_path = create_sample_resume()
        if not resume_path:
            logger.error("Failed to create sample resume.")
            return 1
        
        # Load configuration
        with open('job_scraper_app/config.json', 'r') as f:
            config = json.load(f)
        
        # Update the resume storage path to use the temporary directory
        config["resume_settings"]["storage_path"] = str(resume_path.parent)
        
        # Create database engine
        db_url = config.get('database_url', 'sqlite:///job_listings.db')
        engine = create_engine(db_url)
        
        # Import the required modules
        from job_scraper_app.resume_processor.resume_manager import ResumeManager
        from job_scraper_app.scrapers.scraper_manager import ScraperManager
        
        # Create resume manager
        logger.info("Creating ResumeManager...")
        resume_manager = ResumeManager(config, engine)
        
        # Create scraper manager
        logger.info("Creating ScraperManager...")
        scraper_manager = ScraperManager(config, engine)
        
        # Upload the resume
        logger.info(f"Uploading resume from {resume_path}...")
        resume_id = resume_manager.upload_resume(resume_path, "Sample Resume", True)
        
        if resume_id:
            logger.info(f"Successfully uploaded resume with ID: {resume_id}")
        else:
            logger.error("Failed to upload resume.")
            return 1
        
        # Scrape job listings
        logger.info("Scraping job listings from Remote.co...")
        job_listings = scraper_manager.scrape_site('Remote.co')
        
        if job_listings:
            logger.info(f"Scraped {len(job_listings)} job listings from Remote.co")
            
            # Get the first job listing
            job_listing = job_listings[0]
            logger.info(f"First job listing: {job_listing.title} at {job_listing.company_name}")
            
            # Generate a tailored resume
            logger.info(f"Generating tailored resume for job listing ID: {job_listing.id}...")
            tailored_resume_id = resume_manager.generate_tailored_resume(resume_id, job_listing.id)
            
            if tailored_resume_id:
                logger.info(f"Successfully generated tailored resume with ID: {tailored_resume_id}")
            else:
                logger.error("Failed to generate tailored resume.")
                return 1
            
            # Get tailored resumes
            logger.info(f"Getting tailored resumes for resume ID: {resume_id}...")
            tailored_resumes = resume_manager.get_tailored_resumes(resume_id=resume_id)
            
            if tailored_resumes:
                logger.info(f"Found {len(tailored_resumes)} tailored resumes for resume ID {resume_id}")
                for tr in tailored_resumes:
                    logger.info(f"- {tr.name} (ID: {tr.id})")
            else:
                logger.error(f"No tailored resumes found for resume ID {resume_id}")
                return 1
            
            logger.info("Resume processor test completed successfully!")
            return 0
        else:
            logger.error("Failed to scrape job listings.")
            return 1
    
    except ImportError as e:
        logger.error(f"Import error: {e}")
        logger.error("Make sure all required modules are installed.")
        return 1
    
    except Exception as e:
        logger.error(f"Error: {e}")
        return 1

if __name__ == "__main__":
    sys.exit(test_resume_processor())
